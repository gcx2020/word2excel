import sys
from docx import Document
from xlwt import Workbook



def getClientNames(doc):
  client_names = []

  for paragraph in doc.paragraphs:
    if "Client Name:" in paragraph.text:
      raw = paragraph.text
      client_len = len(raw)
      client_name = raw[12:client_len]
      client_names.append(client_name.lstrip())

  return client_names


def makeHeading(sheet, wb):
  headers = [
    "Client Name",
    "Date", 
    "Contract Type (Original/Additional), Replacement, Informal)",
    "Annual Fee",
    "Comment"
  ]

  for i, header in enumerate(headers):
    sheet.write(0, i, header)

  wb.save("output.xls")

def addTableToSheet(table, sheet, client_name = "N/A", row_offset = 0):
  col_offset = 1  # shifts column count over one for client_name

  # Look at each row, add non-empty cells to the sheet
  for i, row in enumerate(table.rows):
    firstcell = table.cell(i,0)   # first cell of row

    # If firstcell is not empty, consider it
    if firstcell.text != "":
      # If firstcell does no have "Date", add row to sheet
      if not "Date" in firstcell.text:
        # Write client_name to the front of the row
        sheet.write(i + row_offset, 0, client_name)

        # Write cell content to sheet
        for j, cell in enumerate(row.cells):
          
          cell_text = cell.text
          # In the "Comments" cell, replace line break with a comma
          if j == 3:
            cell_text = cell_text.replace('\n', '; ')

          sheet.write(i + row_offset, j + col_offset, cell_text)

      else:
        continue

    else:
      # Consider finished, return new row_offset
      return row_offset + i

  # Return a new row_offset, since rows were added to the sheet
  return row_offset + len(table.rows)

def docTablesToExcel(doc):
  # Create excel file and sheet
  wb = Workbook()
  sheet = wb.add_sheet("Sheet")

  # Get Client Names
  client_names = getClientNames(doc)

  # Add header to excel file
  makeHeading(sheet, wb)

  # Add tables to excel file
  row_offset = 1    # which row to start writing to in the excel sheet
  for i, table in enumerate(doc.tables):
    client_name = client_names[i]
    row_offset = addTableToSheet(table, sheet, client_name, row_offset)

  # Save excel file
  wb.save("output.xls")

def main(argv):
  # Handle Command Line Arguments
  if len(argv) < 2 :
    print("No file given")
    return 1

  # Get docx from file name
  fname = argv[1] 
  doc = Document(fname)

  # Copy tables to a new excel sheet
  docTablesToExcel(doc)

  return 0

# Main Program
if __name__ == "__main__":
  sys.exit( main(sys.argv) )